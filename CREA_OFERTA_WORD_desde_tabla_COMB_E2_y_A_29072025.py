import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor # Importar RGBColor para el color
import os

# --- Reemplazo de placeholders manteniendo formato ---
def reemplazar_placeholders_mejorado(doc, replacements):
    def replace_text_in_paragraph_or_cell(container, key, value):
        full_text = "".join([run.text for run in container.runs])
        if key not in full_text:
            return

        first_run = container.runs[0] if container.runs else None
        font_props = {
            'bold': first_run.font.bold if first_run else None,
            'italic': first_run.font.italic if first_run else None,
            'underline': first_run.font.underline if first_run else None,
            'color': first_run.font.color.rgb if first_run and first_run.font.color else None,
            'size': first_run.font.size if first_run else None,
            'name': first_run.font.name if first_run else None
        }

        new_text = full_text.replace(key, str(value))
        for run in container.runs:
            run.text = ""
        new_run = container.add_run(new_text)

        if first_run:
            if font_props['bold'] is not None: new_run.font.bold = font_props['bold']
            if font_props['italic'] is not None: new_run.font.italic = font_props['italic']
            if font_props['underline'] is not None: new_run.font.underline = font_props['underline']
            if font_props['color'] is not None: new_run.font.color.rgb = font_props['color']
            if font_props['size'] is not None:
                new_run.font.size = font_props['size'] # Simplificado: Pt ya es Pt, si no es None
            if font_props['name'] is not None: new_run.font.name = font_props['name']

    for p in doc.paragraphs:
        for key, value in replacements.items():
            replace_text_in_paragraph_or_cell(p, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_text_in_paragraph_or_cell(p, key, value)

def combinar_documentos_word(documentos_ruta, ruta_salida_combinado):
    if not documentos_ruta:
        messagebox.showwarning("Combinación Vacía", "No hay documentos para combinar.")
        return False

    try:
        documento_combinado = Document(documentos_ruta[0])  # Comienza con 'inicio'
    except Exception as e:
        messagebox.showerror("Error de Combinación", f"No se pudo abrir el documento inicial: {e}")
        return False

    for doc_path in documentos_ruta[1:]:
        try:
            doc_temp = Document(doc_path)

            # NO se agrega salto de página aquí, el control está en la función principal

            for elem in doc_temp.element.body:
                documento_combinado.element.body.append(elem)

        except Exception as e:
            messagebox.showerror("Error al combinar", f"Error al procesar '{doc_path}': {e}")
            return False

    try:
        documento_combinado.save(ruta_salida_combinado)
        print(f"Documento combinado guardado en: {ruta_salida_combinado}")
        return True
    except Exception as e:
        messagebox.showerror("Error al guardar", f"No se pudo guardar el documento combinado: {e}")
        return False

# --- Función principal ---
def generar_documentos_word():
    root = tk.Tk()
    root.withdraw()

    excel_file_path = filedialog.askopenfilename(
        title="Selecciona el archivo Excel con los datos",
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )

    if not excel_file_path:
        messagebox.showwarning("Ningún archivo seleccionado", "No se ha seleccionado ningún archivo Excel.")
        return

    output_directory = os.path.dirname(excel_file_path)

    # Rutas a plantillas
    word_template_insertable = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\modelo_A_AE.docx"
    word_template_centralizado = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\modelo_E2.docx"
    word_template_inicio = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\Of-PRUEBA-00-inicio.docx"
    word_template_oferta = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\Of-PRUEBA-00-final.docx"

    for path in [word_template_insertable, word_template_centralizado, word_template_inicio, word_template_oferta]:
        if not os.path.exists(path):
            messagebox.showerror("Error de Plantilla", f"No se encontró la plantilla:\n{path}")
            return

    try:
        # Obtener número de oferta
        offer_number = os.path.splitext(os.path.basename(excel_file_path))[0].replace("Of-", "", 1)

        # Leer Excel
        headers_df = pd.read_excel(excel_file_path, skiprows=2, nrows=1, header=None)
        column_names = headers_df.iloc[0].tolist()
        data_df = pd.read_excel(excel_file_path, skiprows=4, header=None)
        data_df.columns = column_names

        column_mapping = {
            '{{ITEM}}': 'ITEM',
            '{{CAUDAL}}': 'CAUDAL',
            '{{TEMP}}': 'TEMP',
            '{{PRES}}': 'PRESION',
            '{{SUPF}}': 'SUPERFICIE\nFILTRANTE',
            '{{MODELO}}': 'FILTRO\nMODELO',
            '{{TIPO}}': 'TIPO \nFILTRO',
            '{{KW}}': 'POTENCIA',
            '{{RPM}}': 'VELOCIDAD \nRODETE',
            '{{TRANSMISION}}': 'TRANSMISION',
            '{{---}}': 'WEIGTH',
            '{{PVP}}': 'PVP'
        }

        # Crear documento principal
        doc_principal = Document()
        num_items = 0

        for index, row in data_df.iterrows():
            contador = index + 1
            tipo_filtro = str(row.get('TIPO \nFILTRO', '')).strip().lower()

            if tipo_filtro == "insertable":
                plantilla = word_template_insertable
            elif tipo_filtro == "centralizado":
                plantilla = word_template_centralizado
            else:
                plantilla = word_template_insertable

            try:
                doc_temp = Document(plantilla)
            except Exception as e:
                messagebox.showerror("Error de Plantilla", f"No se pudo abrir '{plantilla}': {e}")
                continue

            replacements = {ph: "" if pd.isna(row.get(col)) else str(row.get(col)) for ph, col in column_mapping.items()}
            replacements.update({
                '{{CONT}}': str(contador),
                '{{UNIDAD}}': "mmca",
                '{{NOF}}': offer_number,
                '{{CONTADOR}}': str(contador)
            })

            reemplazar_placeholders_mejorado(doc_temp, replacements)

            if num_items > 0:
                #doc_principal.add_page_break()
                pass

            for elem in doc_temp.element.body:
                doc_principal.element.body.append(elem)

            num_items += 1

        if num_items == 0:
            messagebox.showwarning("Sin datos", "No se generó ningún ítem. Verifica el archivo Excel.")
            return

        # Guardar documento PVP
        nombre_pvp = f"Of-{offer_number}-00_PVP.docx"
        ruta_pvp = os.path.join(output_directory, nombre_pvp)
        doc_principal.save(ruta_pvp)

        # Combinar con portada y cierre
        documentos_para_combinar = [word_template_inicio, ruta_pvp, word_template_oferta]
        ruta_final = os.path.join(output_directory, f"Of-{offer_number}.docx")

        if combinar_documentos_word(documentos_para_combinar, ruta_final):
            # --- NUEVA LÍNEA: Eliminar el archivo PVP después de una combinación exitosa ---
            try:
                os.remove(ruta_pvp)
                print(f"Archivo temporal '{nombre_pvp}' eliminado correctamente.")
            except Exception as e:
                messagebox.showwarning("Error al eliminar", f"No se pudo eliminar el archivo temporal '{nombre_pvp}':\n{e}")
            # --- FIN NUEVA LÍNEA ---

            messagebox.showinfo("Completado", f"Se generó el documento final:\n{ruta_final}")
        else:
            messagebox.showerror("Error", "No se pudo combinar el documento final.")

    except Exception as e:
        messagebox.showerror("Error Inesperado", f"Ocurrió un error:\n{e}")

if __name__ == "__main__":
    generar_documentos_word()
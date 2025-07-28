import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH # Importar para alineación de párrafo si es necesario
import os # Importamos el módulo os para manejar rutas de archivos

# --- Función mejorada para reemplazar placeholders manteniendo el formato ---
def reemplazar_placeholders_mejorado(doc, replacements):
    """
    Reemplaza todos los marcadores de posición en un documento de Word,
    intentando preservar el formato de los runs existentes.
    """
    # Función auxiliar para reemplazar texto en un párrafo o celda,
    # manteniendo el formato de los runs.
    def replace_text_in_paragraph_or_cell(container, key, value):
        # Concatena el texto de todos los runs para encontrar el placeholder
        full_container_text = "".join([run.text for run in container.runs])

        # Si el placeholder no está en el texto completo, no hay nada que hacer
        if key not in full_container_text:
            return

        # Guarda las propiedades de fuente del primer run del contenedor
        # para intentar heredar el estilo si el contenedor se reconstruye
        first_run_font_properties = None
        if container.runs:
            first_run = container.runs[0]
            first_run_font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'color': first_run.font.color.rgb if first_run.font.color else None,
                'size': first_run.font.size,
                'name': first_run.font.name # Tipo de fuente
            }

        # Reemplaza el placeholder en la cadena completa
        new_full_text = full_container_text.replace(key, str(value))

        # Borra todos los runs existentes en el contenedor (párrafo o celda)
        # Esto es crucial para eliminar el texto antiguo y sus formatos.
        # Recorre la lista en reversa para evitar problemas al eliminar elementos.
        for i in range(len(container.runs) -1, -1, -1):
            container.runs[i]._element.drop_tree() # Elimina el elemento XML del run

        # Crea un nuevo run y asigna el texto reemplazado
        new_run = container.add_run(new_full_text)

        # Intenta aplicar el estilo del primer run original al nuevo run
        if first_run_font_properties:
            if first_run_font_properties['bold'] is not None:
                new_run.font.bold = first_run_font_properties['bold']
            if first_run_font_properties['italic'] is not None:
                new_run.font.italic = first_run_font_properties['italic']
            if first_run_font_properties['underline'] is not None:
                new_run.font.underline = first_run_font_properties['underline']
            if first_run_font_properties['color'] is not None:
                new_run.font.color.rgb = first_run_font_properties['color']
            if first_run_font_properties['size'] is not None:
                new_run.font.size = first_run_font_properties['size']
            if first_run_font_properties['name'] is not None:
                new_run.font.name = first_run_font_properties['name']


    # --- Aplica el reemplazo a párrafos ---
    for p in doc.paragraphs:
        for key, value in replacements.items():
            replace_text_in_paragraph_or_cell(p, key, value)

    # --- Aplica el reemplazo a tablas ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: # Las celdas también contienen párrafos
                    for key, value in replacements.items():
                        replace_text_in_paragraph_or_cell(p, key, value)


def generar_documentos_word():
    """
    Permite al usuario seleccionar un archivo Excel, lee los datos
    y genera un documento Word para cada fila de datos, guardándolos
    en la misma carpeta del Excel seleccionado.
    """
    root = tk.Tk()
    root.withdraw() # Oculta la ventana principal de Tkinter

    excel_file_path = filedialog.askopenfilename(
        title="Selecciona el archivo Excel con los datos",
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )

    if not excel_file_path:
        messagebox.showwarning("Ningún archivo seleccionado", "No se ha seleccionado ningún archivo Excel.")
        return

    # Obtener la ruta del directorio del archivo Excel seleccionado
    output_directory = os.path.dirname(excel_file_path)

    # Definir la plantilla Word (¡AJUSTA ESTA RUTA A LA DE TU PLANTILLA!)
    # Usa una 'r' delante de la cadena para rutas de Windows
    word_template_path = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\modelo_A_AE.docx"

    try:
        # Lee el archivo Excel, saltando las primeras 4 filas (para empezar desde la fila 5)
        # Asegúrate de que la fila 5 de tu Excel contenga los encabezados si quieres usar sus nombres.
        # Si la fila 5 no tiene encabezados, Pandas asignará 'Unnamed: 0', 'Unnamed: 1', etc.
        df = pd.read_excel(excel_file_path, skiprows=4)

        # Mapeo de columnas de Excel a marcadores de posición de Word
        # Si los encabezados de tu Excel son diferentes de 'A', 'D', 'H', etc.,
        # ¡DEBES ACTUALIZAR ESTOS NOMBRES con los nombres reales de tus columnas!
        column_mapping = {
            '{{ITEM}}': 'A',
            '{{CAUDAL}}': 'D',
            '{{TEMP}}': 'H',
            '{{PRES}}': 'G',
            '{{SUPF}}': 'L',
            '{{MODELO}}': 'E',
            '{{TIPO}}': 'F',
            '{{KW}}': 'X',
            '{{RPM}}': 'U',
            '{{---}}': 'AC'
        }

        num_generated_docs = 0
        for index, row_data in df.iterrows():
            contador_fila = index + 1 # El contador empieza en 1 para la primera fila de datos

            # Cargar la plantilla Word para cada nuevo documento
            doc = Document(word_template_path)

            replacements = {}
            for placeholder, col_excel_name in column_mapping.items():
                try:
                    # Acceder al valor de la columna. Convertir a string para reemplazo seguro.
                    # Asegúrate de que el 'col_excel_name' (ej. 'A') exista como encabezado de columna en tu DataFrame.
                    replacements[placeholder] = str(row_data[col_excel_name])
                except KeyError:
                    print(f"Advertencia: Columna '{col_excel_name}' no encontrada en el Excel para el placeholder '{placeholder}' en la fila {contador_fila}. Se usará un valor vacío.")
                    replacements[placeholder] = ""

            # Añadir los valores calculados/fijos
            replacements['{{CONT}}'] = str(contador_fila)
            replacements['{{UNIDAD}}'] = "mmca"
            replacements['{{NOF}}'] = "" # Se mantiene vacío por ahora, ya que no se especificó una columna para ello

            # Reemplazar los marcadores de posición en el documento usando la función mejorada
            reemplazar_placeholders_mejorado(doc, replacements)

            # Construir la ruta completa para el nuevo documento en la carpeta del Excel
            output_file_name = f"documento_generado_{contador_fila}.docx"
            full_output_path = os.path.join(output_directory, output_file_name)

            # Guardar el documento generado
            doc.save(full_output_path)
            num_generated_docs += 1
            print(f"Documento '{full_output_path}' generado con éxito.")

        messagebox.showinfo("Proceso Completado", f"Se han generado {num_generated_docs} documentos Word en:\n{output_directory}")

    except FileNotFoundError:
        messagebox.showerror("Error de archivo", f"La plantilla de Word '{word_template_path}' no se encontró. "
                                                  "Asegúrate de que la ruta sea correcta y el archivo exista.")
    except KeyError as e:
        messagebox.showerror("Error en columnas Excel", f"Una columna esperada no fue encontrada en tu archivo Excel. "
                                                          f"Asegúrate de que los encabezados de la fila 5 del Excel "
                                                          f"(después de saltar las primeras 4 filas) "
                                                          f"coincidan con los nombres de columna esperados en el código ('A', 'D', etc. o tus nombres reales). "
                                                          f"Error: {e}")
    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error inesperado: {e}")

# --- Ejecución del programa ---
if __name__ == "__main__":
    generar_documentos_word()
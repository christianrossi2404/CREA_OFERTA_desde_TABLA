import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION # Importar para saltos de sección
import os

# --- Reemplazo de placeholders manteniendo formato ---
def reemplazar_placeholders_mejorado(doc, replacements):
    def replace_text_in_paragraph_or_cell(container, key, value):
        full_container_text = "".join([run.text for run in container.runs])
        if key not in full_container_text:
            return

        first_run_font_properties = None
        if container.runs:
            first_run = container.runs[0]
            first_run_font_properties = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'color': first_run.font.color.rgb if first_run.font.color else None,
                'size': first_run.font.size,
                'name': first_run.font.name
            }

        new_full_text = full_container_text.replace(key, str(value))
        for run in container.runs:
            run.text = ""
        new_run = container.add_run(new_full_text)

        if first_run_font_properties:
            if first_run_font_properties['bold'] is not None:
                new_run.font.bold = first_run_font_properties['bold']
            if first_run_font_properties['italic'] is not None:
                new_run.font.italic = first_run_font_properties['italic']
            if first_run_font_properties['underline'] is not None:
                new_run.font.underline = first_run_font_properties['underline']
            if first_run_font_properties['color'] is not None:
                new_run.font.color.rgb = first_run_font_properties['color']
            if first_run_font_properties['size'] is not None:
                # Ensure conversion to Pt if not already
                new_run.font.size = (first_run_font_properties['size']
                                      if isinstance(first_run_font_properties['size'], Pt)
                                      else Pt(first_run_font_properties['size']))
            if first_run_font_properties['name'] is not None:
                new_run.font.name = first_run_font_properties['name']

    for p in doc.paragraphs:
        for key, value in replacements.items():
            replace_text_in_paragraph_or_cell(p, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_text_in_paragraph_or_cell(p, key, value)

# --- Combinar documentos Word con saltos de página ---
def combinar_documentos_word(documentos_ruta, ruta_salida_combinado):
    if not documentos_ruta:
        messagebox.showwarning("Combinación Vacía", "No hay documentos para combinar.")
        return False

    try:
        documento_combinado = Document(documentos_ruta[0])
    except FileNotFoundError:
        messagebox.showerror("Error de Combinación", f"El documento base para la combinación '{documentos_ruta[0]}' no se encontró.")
        return False # Indicate failure
    except Exception as e:
        messagebox.showerror("Error de Combinación", f"Error al cargar el primer documento '{documentos_ruta[0]}': {e}")
        return False

    # Iterar desde el segundo documento para añadir saltos de página antes de cada uno
    for i, doc_path in enumerate(documentos_ruta[1:]):
        try:
            doc_temp = Document(doc_path)
            # Añadir un salto de página antes de añadir el contenido del nuevo documento
            # Esto insertará un salto de página antes de 'documentos_generados' y antes de 'word_template_oferta'
            documento_combinado.add_page_break()
            
            for elemento in doc_temp.element.body:
                documento_combinado.element.body.append(elemento)
        except FileNotFoundError:
            messagebox.showwarning("Documento no encontrado", f"El documento '{doc_path}' no se encontró y será omitido en la combinación.")
        except Exception as e:
            messagebox.showerror("Error al combinar", f"Error al procesar '{doc_path}': {e}")
            return False # Indicate failure

    try:
        documento_combinado.save(ruta_salida_combinado)
        print(f"\nDocumento combinado guardado en: {ruta_salida_combinado}")
        return True # Indicate success
    except Exception as e:
        messagebox.showerror("Error al guardar", f"No se pudo guardar el documento combinado en '{ruta_salida_combinado}': {e}")
        return False # Indicate failure

# --- Función principal ---
def generar_documentos_word():
    root = tk.Tk()
    root.withdraw()

    excel_file_path = filedialog.askopenfilename(
        title="Selecciona el archivo Excel con los datos",
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )

    if not excel_file_path:
        messagebox.showwarning("Ningún archivo seleccionado", "No se ha seleccionado ningún archivo Excel.")
        return

    output_directory = os.path.dirname(excel_file_path)

    # Rutas a plantillas
    word_template_insertable = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\modelo_A_AE.docx"
    word_template_centralizado = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\modelo_E2.docx"
    word_template_inicio = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\Of-PRUEBA-00-inicio.docx"
    word_template_oferta = r"C:\Users\Christian Rossi\PRUEBAS\CREA_WORD_OFERTA\PLANTILLAS\Of-PRUEBA-00-final.docx"

    # Validate template paths upfront
    for template_path in [word_template_insertable, word_template_centralizado, word_template_inicio, word_template_oferta]:
        if not os.path.exists(template_path):
            messagebox.showerror("Error de Plantilla", f"La plantilla requerida no se encontró:\n{template_path}\nPor favor, verifica la ruta.")
            return

    try:
        # Obtener número de oferta desde el nombre del archivo
        excel_file_name = os.path.basename(excel_file_path)
        offer_number_raw = os.path.splitext(excel_file_name)[0]
        # More robust way to remove prefix, only if it exists
        offer_number = offer_number_raw.replace("Of-", "", 1)

        # Leer encabezados (fila 3) y datos desde fila 5
        headers_df = pd.read_excel(excel_file_path, skiprows=2, nrows=1, header=None)
        column_names = headers_df.iloc[0].tolist()
        data_df = pd.read_excel(excel_file_path, skiprows=4, header=None)
        data_df.columns = column_names
        df = data_df

        # Mapeo de columnas a marcadores
        column_mapping = {
            '{{ITEM}}': 'ITEM',
            '{{CAUDAL}}': 'CAUDAL',
            '{{TEMP}}': 'TEMP',
            '{{PRES}}': 'PRESION',
            '{{SUPF}}': 'SUPERFICIE\nFILTRANTE',
            '{{MODELO}}': 'FILTRO\nMODELO',
            '{{TIPO}}': 'TIPO \nFILTRO',
            '{{KW}}': 'POTENCIA',
            '{{RPM}}': 'VELOCIDAD \nRODETE',
            '{{TRANSMISION}}': 'TRANSMISION',
            '{{---}}': 'WEIGTH',
            '{{PVP}}': 'PVP'
        }

        num_generated_docs = 0
        documentos_generados = []

        for index, row_data in df.iterrows():
            contador_fila = index + 1

            # Seleccionar plantilla según modelo de filtro
            modelo_filtro = str(row_data.get('TIPO \nFILTRO', '')).strip().lower()
            print(f"Procesando fila {contador_fila}. Tipo de filtro: '{modelo_filtro}'") # Added more detailed print
            if modelo_filtro == "insertable":
                plantilla_usada = word_template_insertable
            elif modelo_filtro == "centralizado":
                plantilla_usada = word_template_centralizado
            else:
                plantilla_usada = word_template_insertable  # por defecto
                print(f"Advertencia: Tipo de filtro '{modelo_filtro}' no reconocido. Usando plantilla por defecto: {plantilla_usada}")

            try:
                doc = Document(plantilla_usada)
            except FileNotFoundError:
                messagebox.showerror("Error de Plantilla", f"La plantilla '{plantilla_usada}' no se encontró al procesar la fila {contador_fila}. Saltando esta fila.")
                continue # Skip to the next row if template is missing

            replacements = {}
            for placeholder, col_excel_name in column_mapping.items():
                valor = row_data.get(col_excel_name)
                replacements[placeholder] = "" if pd.isna(valor) else str(valor)

            replacements['{{CONT}}'] = str(contador_fila)
            replacements['{{UNIDAD}}'] = "mmca"
            replacements['{{NOF}}'] = offer_number
            replacements['{{CONTADOR}}'] = str(contador_fila)

            reemplazar_placeholders_mejorado(doc, replacements)

            output_file_name = f"documento_generado_{contador_fila}.docx"
            full_output_path = os.path.join(output_directory, output_file_name)
            doc.save(full_output_path)

            documentos_generados.append(full_output_path)
            num_generated_docs += 1
            print(f"Documento '{full_output_path}' generado con éxito.")

        # Combinar todos los documentos generados
        if documentos_generados:
            # La lista de documentos a combinar ahora incluirá saltos de página entre ellos
            documentos_para_combinar = [word_template_inicio] + documentos_generados + [word_template_oferta]

            ruta_doc_combinado = os.path.join(output_directory, f"Of-{offer_number}_COMBINADO.docx")
            
            if combinar_documentos_word(documentos_para_combinar, ruta_doc_combinado):
                messagebox.showinfo("Proceso Completado", f"Se han generado {num_generated_docs} documentos intermedios y el documento combinado se guardó en:\n{ruta_doc_combinado}")

                # Eliminar todos los archivos que empiecen por "documento_generado_"
                print("\nIniciando limpieza de archivos temporales...")
                for archivo in os.listdir(output_directory):
                    if archivo.startswith("documento_generado_") and archivo.endswith(".docx"):
                        try:
                            os.remove(os.path.join(output_directory, archivo))
                            print(f"Archivo temporal eliminado: {archivo}")
                        except Exception as e:
                            print(f"No se pudo eliminar {archivo}: {e}")
                print("Limpieza completada.")
            else:
                messagebox.showerror("Proceso Incompleto", "Hubo un error al combinar los documentos. Los archivos intermedios pueden no haber sido eliminados.")


    except Exception as e:
        messagebox.showerror("Error Inesperado", f"Ha ocurrido un error inesperado durante el procesamiento: {e}")
        # Optionally, log the full traceback for debugging
        # import traceback
        # print(traceback.format_exc())

if __name__ == "__main__":
    generar_documentos_word()